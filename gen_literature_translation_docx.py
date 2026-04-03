# -*- coding: utf-8 -*-
"""生成文献翻译 Word 文档（完整版：含全文英文与中文译文）"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from paper_full_content import (
    EN_TITLE, EN_AUTHORS, EN_SOURCE, EN_SECTIONS,
    CN_TITLE, CN_AUTHORS, CN_SOURCE, CN_SECTIONS,
)

def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 0 else 14)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(12)
    return p

def add_para(doc, text, font_size=11):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(6)
    return p

def add_para_cn(doc, text, font_size=11):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(6)
    return p

def add_section_heading(doc, text, is_cn=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman' if not is_cn else '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def main():
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # 封面
    add_heading_cn(doc, "毕业设计（论文）文献翻译", 0)
    add_heading_cn(doc, "—— 网络协议模糊测试与种子调度", 0)
    doc.add_paragraph()
    doc.add_paragraph()

    # ========== 一、英文原文 ==========
    h1 = doc.add_heading("一、英文原文", level=1)
    h1.runs[0].font.name = '宋体'
    h1.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    r = p_title.add_run(EN_TITLE)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'
    p_title.paragraph_format.space_after = Pt(6)
    add_para(doc, EN_AUTHORS, 10)
    add_para(doc, EN_SOURCE, 9)
    doc.add_paragraph()

    for sec_title, paragraphs in EN_SECTIONS:
        add_section_heading(doc, sec_title, is_cn=False)
        for para_text in paragraphs:
            for block in para_text.split("\n\n"):
                block = block.strip()
                if block:
                    add_para(doc, block, 11)
        doc.add_paragraph()

    doc.add_paragraph()
    doc.add_paragraph()

    # ========== 二、中文译文 ==========
    h2 = doc.add_heading("二、中文译文", level=1)
    h2.runs[0].font.name = '宋体'
    h2.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()

    p_cn_title = doc.add_paragraph()
    r2 = p_cn_title.add_run(CN_TITLE)
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.name = '宋体'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p_cn_title.paragraph_format.space_after = Pt(6)
    add_para_cn(doc, CN_AUTHORS, 10)
    add_para_cn(doc, CN_SOURCE, 9)
    doc.add_paragraph()

    for sec_title, paragraphs in CN_SECTIONS:
        add_section_heading(doc, sec_title, is_cn=True)
        for para_text in paragraphs:
            for block in para_text.split("\n\n"):
                block = block.strip()
                if block:
                    add_para_cn(doc, block, 11)
        doc.add_paragraph()

    out_path = r"f:\学习\学校\毕业设计\代码\文献翻译-基于强化学习的网络协议模糊测试多阶段种子调度-完整版.docx"
    doc.save(out_path)
    print("已生成（完整版）:", out_path)

if __name__ == "__main__":
    main()
